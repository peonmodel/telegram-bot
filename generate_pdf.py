from docx import Document
from docx.shared import Inches
import docx2pdf

# # sample input
# result =[
#     {"key": "project.title", "value": "ABC water Project"}, 
#     {"key": "project.ref", "value": "A1234-12345-2022"}, 
#     {"key": "attendees.0.name", "value": "Ong Jun Yong"}, 
#     {"key": "attendees.0.designation", "value": "Project Manager"}, 
#     {"key": "attendees.0.photo", "value": "AgACAgUAAxkBAAMPYxXl1WkIQ8KVgH9hVSZyJXyJ-HwAAliyMRtPtbBU9w5S0FDMR5YBAAMCAANtAAMpBA.jpg"}, 
#     {"key": "attendees.1.name", "value": "Lee Xuan Yen"}, 
#     {"key": "attendees.1.designation", "value": "Modeller"}, 
#     {"key": "attendees.1.photo", "value": "AgACAgUAAxkBAAMYYxXmBHBTFdbLJnuFU7JVYqW1LfQAAlqyMRtPtbBUGRgemsAcDZoBAAMCAANtAAMpBA.jpg"}, 
#     {"key": "attendees.2.name", "value": "Chen Shujun"}, 
#     {"key": "attendees.2.designation", "value": "executor"}, 
#     {"key": "attendees.2.photo", "value": "AgACAgUAAxkBAAMhYxXmKzbf289vfiZYO7-GlEfBPP8AAluyMRtPtbBUhvoHel_mK5cBAAMCAANtAAMpBA.jpg"}, 
#     {"key": "issues.0.image", "value": "AgACAgUAAxkBAAMmYxXmRxA4hzm5Ma8G5Brq9zoYqJsAAlyyMRtPtbBUZm-zTb7uLsYBAAMCAANtAAMpBA.jpg"}, 
#     {"key": "issues.0.caption", "value": "Wondering what is this?"}, 
#     {"key": "issues.0.raisedBy", "value": "Xuan Yen"}, 
#     {"key": "issues.0.raisedAt", "value": "Telegram"}, 
#     {"key": "issues.1.image", "value": "AgACAgUAAxkBAAMxYxXmeQJD9VedG4KIm5sPGBSdJOkAAliyMRtPtbBU9w5S0FDMR5YBAAMCAANtAAMpBA.jpg"}, 
#     {"key": "issues.1.caption", "value": "Can't stop from laughing"}, 
#     {"key": "issues.1.raisedBy", "value": "Xuan Yen"},
#     {"key": "issues.1.raisedAt", "value": "home"}, 
#     {"key": "tags.0", "value": "Shujun"}
# ]

async def generate_pdf(template: dict, record: dict):
    project_title = ""
    project_ref = ""
    attendees_name =[]
    attendees_designation =[]
    attendees_photo =[]
    issues_image =[]
    issues_caption =[]
    issues_raisedBy =[]
    issues_raisedAt =[]


    i = 0
    for input in record:
        if record [i]["key"] == "project.title":
            project_title = record [i]["value"]
        elif record [i]["key"] == "project.ref":
            project_ref = record [i]["value"]
        elif (record [i]["key"].split("."))[0] == "attendees" and (record [i]["key"].split("."))[2] == "name":
            attendees_name.append(record [i]["value"])
        elif (record [i]["key"].split("."))[0] == "attendees" and (record [i]["key"].split("."))[2] == "designation":
            attendees_designation.append(record [i]["value"])
        elif (record [i]["key"].split("."))[0] == "attendees" and (record [i]["key"].split("."))[2] == "photo":
            attendees_photo.append(record [i]["value"])
        elif (record [i]["key"].split("."))[0] == "issues" and (record [i]["key"].split("."))[2] == "image":
            issues_image.append(record [i]["value"])
        elif (record [i]["key"].split("."))[0] == "issues" and (record [i]["key"].split("."))[2] == "caption":
            issues_caption.append(record [i]["value"])
        elif (record [i]["key"].split("."))[0] == "issues" and (record [i]["key"].split("."))[2] == "raisedBy":
            issues_raisedBy.append(record [i]["value"])
        elif (record [i]["key"].split("."))[0] == "issues" and (record [i]["key"].split("."))[2] == "raisedAt":
            issues_raisedAt.append(record [i]["value"])
        i = i + 1

    filename = template['document_title'] + "_" + project_title

    document = Document()

    document.add_heading('Inspection Report', 0)

    document.add_paragraph('Project title:\t\t {}'.format(project_title))
    document.add_paragraph('Project reference:\t {}'.format(project_ref))

    document.add_paragraph('Attendees')

    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S/N'
    hdr_cells[1].text = 'Attendee Name'
    hdr_cells[2].text = 'Designation'
    hdr_cells[3].text = 'Photo'

    sn_attendee = 1
    for name, designation, photo in zip (attendees_name, attendees_designation, attendees_photo):
        row_cell = table.add_row().cells
        row_cell[0].text = str(sn_attendee)
        row_cell[1].text = name
        row_cell[2].text = designation
        
        attendee_cell = table.rows[sn_attendee].cells[3]
        paragraph = attendee_cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture('./telegram input/'+photo, width=Inches(1.25))
        sn_attendee = sn_attendee + 1
      
    document.add_paragraph('\n')
    document.add_paragraph('Issues observed')

    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S/N'
    hdr_cells[1].text = 'Issues'
    hdr_cells[2].text = 'Description'
    hdr_cells[3].text = 'Raised by'
    hdr_cells[4].text = 'Raised at'

    sn = 1
    for image, caption, raisedBy, raisedAt in zip (issues_image, issues_caption, issues_raisedBy, issues_raisedAt):
            row_cell = table.add_row().cells
            row_cell[0].text = str(sn)
                
            issue_cell = table.rows[sn].cells[1]
            paragraph = issue_cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture('./telegram input/'+image, width=Inches(2.25))
                
            row_cell[2].text = caption
            row_cell[3].text = raisedBy
            row_cell[4].text = raisedAt
            sn = sn + 1

    document.add_paragraph('\n')

    document.save('./generated pdf/{}.docx'.format(filename))
    
    word_file ='./generated pdf/{}.docx'.format(filename)
    pdf_file = './generated pdf/{}.pdf'.format(filename)
    with open(pdf_file, "wb") as f:
       pass
    docx2pdf.convert(word_file, pdf_file)
    return filename
    
# generate_pdf({}, result)
