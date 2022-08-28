from docx import Document
from docx.shared import Inches

#Sample Input needed
#project information
project_ref = "A1234-00001-2021"
project_title = "MMT Kranji Loop"
date_inspection = "07/08/2022"
time_inspection = "10:35"
inspector = "Ms Lee"
designation = "Site supervisor"
reviewer = "Mr Ong"
approver = "Ms Chen"
filename = "inspection report_2"

#Inspection information
photos = ["photo 1.png","photo 2.png", "photo 3.png"]
description = ["desc 1","desc 2","desc 3"]
remarks =["remark 1","remark 2","remark 3"]

qty=[]
sn = 1

for x in photos:
    qty.append(sn)
    sn = sn + 1

document = Document()

document.add_heading('Inspection Checklist', 0)

document.add_paragraph('Project reference:\t {}'.format(project_ref))
document.add_paragraph('Project title:\t\t {}'.format(project_title))
document.add_paragraph('Date:\t\t\t {}'.format(date_inspection))
document.add_paragraph('Time:\t\t\t {}'.format(time_inspection))
document.add_paragraph('Inspector name:\t {}'.format(inspector))
document.add_paragraph('Designation:\t\t {}'.format(designation))
document.add_paragraph('\n')

#document.add_picture('photo 1.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'S/N'
hdr_cells[1].text = 'Photos'
hdr_cells[2].text = 'Description'
hdr_cells[3].text = 'Remarks'

for qty, photo, desc, remark in zip(qty, photos, description, remarks):
    row_cell = table.add_row().cells
    row_cell[0].text = str(qty)
    
    photo_cell = table.rows[qty].cells[1]
    paragraph = photo_cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(photos[qty-1], width=Inches(2.25))
    
    row_cell[2].text = desc
    row_cell[3].text = remark

document.add_paragraph('\n')

document.add_paragraph('Submitted by:\t {}'.format(inspector))

document.save('{}.docx'.format(filename))


import docx2pdf 
word_file ='{}.docx'.format(filename)
pdf_file = '{}.pdf'.format(filename)
with open(pdf_file, "wb") as f:
    pass
docx2pdf.convert(word_file, pdf_file)


